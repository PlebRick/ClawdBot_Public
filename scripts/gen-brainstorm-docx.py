#!/usr/bin/env python3
"""Generate a Bible Brainstorm .docx from markdown sections."""
import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def build_doc(passage_text, outline_text, summary_text, appendix_text, output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Bible Brainstorm: Ephesians 2:1-10', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Generated: January 30, 2026 | Research material for sermon-writer', italic=True, size=10)
    doc.add_paragraph()
    
    # Section 1: Passage
    add_heading(doc, 'The Passage — Ephesians 2:1-10 (ESV)', level=1)
    add_para(doc, passage_text, italic=True)
    doc.add_page_break()
    
    # Section 2: Outline
    add_heading(doc, 'Working Outline — "But God"', level=1)
    for line in outline_text.strip().split('\n'):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=2)
        elif line.startswith('**') and line.endswith('**'):
            add_para(doc, line.strip('*'), bold=True)
        elif line.startswith('- **Key Phrase:**'):
            p = doc.add_paragraph()
            run = p.add_run('Key Phrase: ')
            run.bold = True
            run.font.size = Pt(11)
            rest = line.replace('- **Key Phrase:**', '').strip().strip('"')
            run2 = p.add_run(f'"{rest}"')
            run2.italic = True
            run2.font.size = Pt(11)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('  - '):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')
        else:
            add_para(doc, line)
    doc.add_page_break()
    
    # Section 3: Summary Draft
    add_heading(doc, 'Summary Draft', level=1)
    add_para(doc, 'Research material — not final sermon. To be fed into sermon-writer with commentaries and other sources.', italic=True, size=10)
    doc.add_paragraph()
    for line in summary_text.strip().split('\n'):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            add_para(doc, line.strip('*'), italic=True)
        else:
            add_para(doc, line)
    doc.add_page_break()
    
    # Section 4: Research Appendix
    add_heading(doc, 'Research Appendix', level=1)
    for line in appendix_text.strip().split('\n'):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('- **'):
            # Bold lead item
            parts = line[2:].split('**')
            p = doc.add_paragraph(style='List Bullet')
            if len(parts) >= 3:
                run = p.add_run(parts[1])
                run.bold = True
                run.font.size = Pt(11)
                run2 = p.add_run(parts[2])
                run2.font.size = Pt(11)
            else:
                p.add_run(line[2:]).font.size = Pt(11)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.size = Pt(10)
        else:
            add_para(doc, line)
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")

if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "output.docx"
    
    # Read sections from stdin or files
    passage_file = sys.argv[2] if len(sys.argv) > 2 else None
    outline_file = sys.argv[3] if len(sys.argv) > 3 else None
    summary_file = sys.argv[4] if len(sys.argv) > 4 else None
    appendix_file = sys.argv[5] if len(sys.argv) > 5 else None
    
    def read_file(path):
        with open(path) as f:
            return f.read()
    
    passage = read_file(passage_file) if passage_file else ""
    outline = read_file(outline_file) if outline_file else ""
    summary = read_file(summary_file) if summary_file else ""
    appendix = read_file(appendix_file) if appendix_file else ""
    
    build_doc(passage, outline, summary, appendix, output)
