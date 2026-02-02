#!/usr/bin/env python3
"""
Generate a liturgy handout .docx for St. Peter's Stone Church.

Usage: python3 generate-liturgy.py --json '<json_payload>' --output /path/to/output.docx

JSON payload:
{
  "passage": "Luke 6:27-38",
  "title": "Love Your Enemies",
  "date_display": "February 23, 2025",
  "sunday_name": "Seventh Sunday after the Epiphany",
  "year": "C",
  "readings": [
    {"label": "Old Testament", "citation": "Genesis 45:3-11, 15", "reader": "Liturgist"},
    {"label": "Psalm", "citation": "Psalm 37:1-11, 39-40", "reader": "Liturgist"},
    {"label": "New Testament", "citation": "1 Corinthians 15:35-38, 42-50", "reader": "Liturgist"},
    {"label": "Gospel", "citation": "Luke 6:27-38", "reader": "Chaplain"}
  ],
  "call_to_worship": [
    {"speaker": "Leader", "text": "Love your enemies..."},
    {"speaker": "People", "text": "Bless those who curse you..."}
  ],
  "benediction": {"text": "May the God of hope...", "citation": "Romans 15:13, ESV"},
  "theme": "Love Your Enemies"
}
"""

import json
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_text(doc, text, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_body_text(doc, text, bold=False, italic=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    # Add some space before
    p.paragraph_format.space_before = Pt(12)
    return p

def add_bullet(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        run = p.add_run(f"• {bold_prefix}")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
        if italic:
            run.italic = True
    else:
        run = p.add_run(f"• {text}")
        run.font.size = Pt(11)
        if italic:
            run.italic = True
    return p

def generate(data, output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title block
    add_heading_text(doc, f"{data['passage']} {data['title']}", size=16)
    add_heading_text(doc, f"Date: {data['date_display']}", size=11, bold=False)
    add_heading_text(doc, f"{data['sunday_name']}, Year {data['year']}", size=11, bold=False)
    add_heading_text(doc, f"Theme: {data['theme']}", size=11, bold=True)
    
    # Introduction
    add_section_header(doc, "Introduction")
    add_bullet(doc, "Ringing of the bell")
    add_bullet(doc, "Prelude – music")
    add_bullet(doc, "Light enters the sanctuary")
    add_bullet(doc, "Welcome")
    add_bullet(doc, "Announcements")
    add_bullet(doc, "Share peace and greetings")
    
    # Worship
    add_section_header(doc, "Worship")
    add_bullet(doc, "Hymn: ", bold_prefix="Hymn: ", italic=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• ")
    run.font.size = Pt(11)
    run = p.add_run("Call to Worship: ")
    run.bold = True
    run.font.size = Pt(11)
    # Add psalm attribution if available
    psalm_cite = None
    for r in data.get('readings', []):
        if 'Psalm' in r.get('label', ''):
            psalm_cite = r['citation']
            break
    if psalm_cite:
        run2 = p.add_run(f'(Based on {psalm_cite})')
        run2.italic = True
        run2.font.size = Pt(11)
    run.bold = True
    run.font.size = Pt(11)
    
    # Call to Worship responsive reading
    for line in data['call_to_worship']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"{line['speaker']}: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(line['text'])
        run.font.size = Pt(11)
    
    # Opening Prayer
    add_bullet(doc, "Opening Prayer")
    add_body_text(doc, "(Chaplain: extemporaneous)", italic=True, indent=0.5)
    
    # The Word
    add_section_header(doc, "The Word: Scripture Reading")
    for reading in data['readings']:
        add_bullet(doc, f" {reading['citation']} ({reading['reader']})", bold_prefix=f"{reading['label']}: ")
    
    # Sermon
    add_bullet(doc, f" {data['theme']}", bold_prefix="Sermon: ")
    
    # Moment of silence
    add_body_text(doc, "Moment of silence for reflection and response accompanied by music", italic=True, size=11)
    add_body_text(doc, "(Hymn to be chosen)", italic=True, indent=0.25)
    
    # Prayer of Confession - Apostles' Creed
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• ")
    run.font.size = Pt(11)
    run = p.add_run("Prayer of Confession")
    run.bold = True
    run.font.size = Pt(11)
    
    creed = ("I believe in God, the Father almighty, creator of heaven and earth. "
             "I believe in Jesus Christ, his only Son, our Lord, who was conceived by the Holy Spirit "
             "and born of the virgin Mary. He suffered under Pontius Pilate, was crucified, died, and was buried... "
             "The third day he rose again from the dead. He ascended to heaven and is seated at the right hand "
             "of God the Father almighty. From there he will come to judge the living and the dead.")
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run("All: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(creed)
    run.font.size = Pt(11)
    
    add_body_text(doc, "(taken from the Apostles Creed)", italic=True, indent=0.5, size=10)
    
    # Prayer and Offering
    add_section_header(doc, "Prayer and Offering")
    add_bullet(doc, "Hymn: ", bold_prefix="Hymn: ", italic=True)
    add_bullet(doc, "Prayers of the faithful (Prayer requests of the congregation)")
    
    # Lord's Prayer
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• ")
    run.font.size = Pt(11)
    run = p.add_run("The Lord's prayer")
    run.bold = True
    run.font.size = Pt(11)
    
    lords_prayer = ("Our Father which art in heaven, Hallowed be thy name. "
                    "Thy kingdom come, Thy will be done in earth, as it is in heaven. "
                    "Give us this day our daily bread. And forgive us our debts, as we forgive our debtors. "
                    "And lead us not into temptation, but deliver us from evil: "
                    "For thine is the kingdom, and the power, and the glory, for ever. Amen.")
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(lords_prayer)
    run.font.size = Pt(11)
    
    add_body_text(doc, "(Matthew 6:9-13)", italic=True, indent=0.5, size=10)
    
    # Offering
    add_bullet(doc, "Offering (Prayer of blessing over offering)")
    
    # Closing
    add_section_header(doc, "Closing")
    add_bullet(doc, "Closing hymn: ", bold_prefix="Closing hymn: ", italic=True)
    
    # Benediction
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• ")
    run.font.size = Pt(11)
    run = p.add_run("Benediction: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(f'"{data["benediction"]["text"]}" ({data["benediction"]["citation"]})')
    run.font.size = Pt(11)
    
    # Amen
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• ")
    run.font.size = Pt(11)
    run = p.add_run("All: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("Amen!")
    run.font.size = Pt(11)
    
    # Postlude
    add_bullet(doc, "Postlude - music")
    
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--json', required=True, help='JSON payload')
    parser.add_argument('--output', required=True, help='Output .docx path')
    args = parser.parse_args()
    
    data = json.loads(args.json)
    generate(data, args.output)
